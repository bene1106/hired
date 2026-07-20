"""Build the submission .docx from PROJECT_DOC.md.

Content is transformed, never retyped, so the .docx cannot drift from the
verified source. Sections are remapped onto the numbering used by the course's
Requirements template (0. Cover .. 10. Appendices).

Run:  uv run --with python-docx python build_docx.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2])

# Requirements-template section -> (title, [source section numbers], demote_after_first)
# The first source keeps the target heading; later ones are demoted into it.
MAPPING = [
    ("1. Problem Statement & Vision", [1]),
    ("2. Scope", [2]),
    ("3. User Stories", [3]),
    ("4. System Overview", [4, 6]),          # + UI/UX Design
    ("5. Tech Stack & External Services", [5, 7, 8]),  # + Ethics + External Services
    ("6. Project Plan & Timeline", [9]),
    ("7. Roles & Responsibilities", [10]),
    ("8. Risk Register", [11]),
    ("9. Evaluation & Demo Plan", [12]),
    ("10. Appendices", [13]),
]


def split_sections(md: str) -> dict[int, tuple[str, list[str]]]:
    """Return {section_number: (title, body_lines)} for each '## N. Title'."""
    out: dict[int, tuple[str, list[str]]] = {}
    cur: int | None = None
    title = ""
    buf: list[str] = []
    for line in md.splitlines():
        m = re.match(r"^##\s+(\d+)\.\s+(.*)$", line)
        if m:
            if cur is not None:
                out[cur] = (title, buf)
            cur = int(m.group(1))
            title = m.group(2).strip()
            buf = []
            continue
        if cur is not None:
            buf.append(line)
    if cur is not None:
        out[cur] = (title, buf)
    return out


# --------------------------------------------------------------------------
# Inline formatting: **bold**, `code`, *italic*
# --------------------------------------------------------------------------
TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(par, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links -> label only
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        else:
            par.add_run(piece)


def emit(doc: Document, lines: list[str], demote: int = 0) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1)) - 1 + demote  # '##'->1
            doc.add_heading(re.sub(r"^\d+(\.\d+)*\.?\s*", "", m.group(2)), min(level, 4))
            i += 1
            continue

        # Images
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if m:
            path = (SRC.parent / m.group(2)).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                print(f"  WARNING: image not found: {path}")
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                width = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=width)
                tbl.style = "Light Grid Accent 1"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ridx, row in enumerate(rows):
                    cells = tbl.add_row().cells
                    for cidx in range(width):
                        cell = cells[cidx]
                        cell.text = ""
                        par = cell.paragraphs[0]
                        add_runs(par, row[cidx] if cidx < len(row) else "")
                        for run in par.runs:
                            run.font.size = Pt(8.5)
                            if ridx == 0:
                                run.bold = True
                doc.add_paragraph()
            continue

        # Fenced code
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            run = par.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            add_runs(par, " ".join(quote))
            for run in par.runs:
                run.italic = True
            continue

        # Bullets / numbered
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            body = [m.group(1)]
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip():
                body.append(lines[i].strip())
                i += 1
            add_runs(doc.add_paragraph(style="List Bullet"), " ".join(body))
            continue

        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            add_runs(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        # Paragraph (join wrapped lines)
        body = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(r"^([-*#>|]|\d+\.|```)", nxt):
                break
            body.append(nxt)
            i += 1
        add_runs(doc.add_paragraph(), " ".join(body))


def main() -> int:
    md = SRC.read_text(encoding="utf8")
    sections = split_sections(md)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ---- 0. Cover -------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Hired.")
    run.bold = True
    run.font.size = Pt(40)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("AI-Powered Career Planning Agent")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("A local-first desktop application. The user's data never leaves their machine.")
    run.italic = True

    for _ in range(2):
        doc.add_paragraph()

    for label, value in [
        ("Project", "Hired. — Local-First AI Career Agent"),
        ("Team", "Anna Vegera · Benedict Herrnleben · Eren Kocadag · Muhammad Kaleem Ullah"),
        ("Course", "SWP SS 2026 — Chat, Search and Summaries: Smarter Apps with LLMs"),
        ("Version", "v1.0 — describes the shipped system as of release v0.5.0"),
        ("Date", "20 July 2026 (first drafted 23 April 2026)"),
        ("Repository", "github.com/bene1106/hired"),
    ]:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(f"{label}: ")
        run.bold = True
        par.add_run(value)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Where something was planned but not built, this document says so rather "
        "than describing the plan as if it were the product."
    )
    run.italic = True
    run.font.size = Pt(9.5)

    doc.add_page_break()

    # ---- 1..10 ----------------------------------------------------------
    for target_title, src_nums in MAPPING:
        doc.add_heading(target_title, 1)
        for idx, num in enumerate(src_nums):
            if num not in sections:
                print(f"  WARNING: source section {num} not found")
                continue
            src_title, body = sections[num]
            if idx > 0:
                doc.add_heading(src_title, 2)
                emit(doc, body, demote=1)
            else:
                emit(doc, body, demote=0)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  sections mapped: {len(MAPPING)}  source sections used: {sum(len(s) for _, s in MAPPING)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
